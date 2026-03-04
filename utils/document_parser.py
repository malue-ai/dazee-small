"""
统一文档解析器 — Document Parser

三级降级链：
    Unstructured API (云端, 表格+结构+OCR)
        ↓ API 不可用
    pdfplumber (本地, 表格+文本)
        ↓ 未安装
    PyPDF2 (本地, 仅文本)

职责：
    将 PDF/DOCX 解析为结构化 ParsedDocument，供 file_processor 和 Skill 使用。
"""

import os
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

from logger import get_logger

logger = get_logger("document_parser")


# ==================== 数据模型 ====================


class ElementType(str, Enum):
    TITLE = "Title"
    NARRATIVE_TEXT = "NarrativeText"
    TABLE = "Table"
    LIST_ITEM = "ListItem"
    HEADER = "Header"
    FOOTER = "Footer"
    PAGE_BREAK = "PageBreak"
    IMAGE = "Image"
    FORMULA = "Formula"
    UNCATEGORIZED = "UncategorizedText"


@dataclass
class Element:
    type: ElementType
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """文档解析结果"""
    elements: List[Element] = field(default_factory=list)
    markdown: str = ""
    page_count: int = 0
    table_count: int = 0
    parser_used: str = ""

    @property
    def summary(self) -> str:
        """生成适合注入上下文的摘要"""
        titles = [
            e.text for e in self.elements
            if e.type == ElementType.TITLE
        ][:15]
        toc = "\n".join(f"  - {t}" for t in titles) if titles else "  (未检测到章节标题)"
        preview = self.markdown[:500] if self.markdown else ""
        return (
            f"页数: {self.page_count}, 表格: {self.table_count}, "
            f"解析器: {self.parser_used}\n"
            f"章节目录:\n{toc}\n"
            f"内容预览:\n{preview}"
        )


# ==================== 解析器实现 ====================


class DocumentParser:
    """统一文档解析器：Unstructured API (主) + pdfplumber (降级) + PyPDF2 (兜底)"""

    def __init__(self) -> None:
        # API Key 可来自：环境变量、实例 .env、或对话 HITL 持久化到本地 SQLite 后注入的 os.environ
        self._api_key: Optional[str] = os.getenv("UNSTRUCTURED_API_KEY")
        self._api_url: str = os.getenv(
            "UNSTRUCTURED_API_URL",
            "https://api.unstructuredapp.io",
        )

    def _unstructured_available(self) -> bool:
        if not self._api_key:
            logger.info("UNSTRUCTURED_API_KEY 未配置，将使用本地解析 (pdfplumber/PyPDF2)")
            return False
        try:
            import unstructured_client  # noqa: F401
            return True
        except ImportError:
            logger.warning("unstructured_client 未安装，降级到本地解析。安装: pip install unstructured-client")
            return False

    def _pdfplumber_available(self) -> bool:
        try:
            import pdfplumber  # noqa: F401
            return True
        except ImportError:
            return False

    # -------------------- 统一入口 --------------------

    async def parse(
        self,
        file_path: str,
        *,
        strategy: str = "hi_res",
        page_range: Optional[List[int]] = None,
        chunking: bool = False,
        chunk_max_chars: int = 4000,
    ) -> ParsedDocument:
        """
        解析 PDF/DOCX 文件，返回结构化结果。

        Args:
            file_path: 文件本地路径
            strategy: Unstructured 策略 (hi_res / fast / auto)
            page_range: PDF 页码范围 [start, end]，只解析指定区间
            chunking: 是否启用 by_title 分块（按章节标题自动切分）
            chunk_max_chars: chunking 时每 chunk 最大字符数
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return await self._parse_docx(path)

        if suffix != ".pdf":
            logger.warning(f"不支持的文件类型: {suffix}，尝试按 PDF 解析")

        if self._unstructured_available():
            try:
                return await self._parse_with_unstructured(
                    path, strategy,
                    page_range=page_range,
                    chunking=chunking,
                    chunk_max_chars=chunk_max_chars,
                )
            except Exception as e:
                logger.warning(f"Unstructured API 失败，降级到本地解析: {e}")

        if self._pdfplumber_available():
            try:
                return await self._parse_with_pdfplumber(path)
            except Exception as e:
                logger.warning(f"pdfplumber 失败，降级到 PyPDF2: {e}")

        return await self._parse_with_pypdf2(path)

    # -------------------- Unstructured API --------------------

    async def _parse_with_unstructured(
        self,
        path: Path,
        strategy: str,
        *,
        page_range: Optional[List[int]] = None,
        chunking: bool = False,
        chunk_max_chars: int = 4000,
    ) -> ParsedDocument:
        from unstructured_client import UnstructuredClient
        from unstructured_client.models.shared import Files, PartitionParameters
        from unstructured_client.models.operations import PartitionRequest

        client = UnstructuredClient(
            api_key_auth=self._api_key,
            server_url=self._api_url,
        )

        with open(path, "rb") as f:
            file_content = f.read()

        kw: Dict[str, Any] = dict(
            files=Files(content=file_content, file_name=path.name),
            strategy=strategy,
            languages=["eng"],
            pdf_infer_table_structure=True,
            split_pdf_page=True,
            split_pdf_allow_failed=True,
            split_pdf_concurrency_level=15,
            include_page_breaks=True,
        )
        if page_range and len(page_range) == 2:
            kw["split_pdf_page_range"] = page_range
        if chunking:
            kw["chunking_strategy"] = "by_title"
            kw["max_characters"] = chunk_max_chars
            kw["new_after_n_chars"] = int(chunk_max_chars * 0.75)
            kw["combine_under_n_chars"] = 500
            kw["multipage_sections"] = True

        params = PartitionParameters(**kw)
        req = PartitionRequest(partition_parameters=params)
        resp = client.general.partition(request=req)

        raw_elements = resp.elements or []

        elements: List[Element] = []
        table_count = 0
        pages_seen: set[int] = set()

        for item in raw_elements:
            item_dict = item if isinstance(item, dict) else (
                item.model_dump() if hasattr(item, "model_dump") else
                item.__dict__ if hasattr(item, "__dict__") else {"text": str(item)}
            )
            etype = _map_unstructured_type(item_dict.get("type", ""))
            text = item_dict.get("text", "")
            raw_meta = item_dict.get("metadata", {})
            if isinstance(raw_meta, str):
                meta: Dict[str, Any] = {}
            elif hasattr(raw_meta, "model_dump"):
                meta = raw_meta.model_dump()  # type: ignore[union-attr]
            elif hasattr(raw_meta, "__dict__"):
                meta = dict(raw_meta.__dict__)
            else:
                meta = raw_meta if isinstance(raw_meta, dict) else {}
            page = meta.get("page_number")
            if page:
                pages_seen.add(page)
            if etype == ElementType.TABLE:
                table_count += 1
                html = meta.get("text_as_html", "")
                if html:
                    text = _html_table_to_markdown(html) or text

            elements.append(Element(
                type=etype,
                text=text,
                metadata={"page_number": page},
            ))

        md = _elements_to_markdown(elements)
        return ParsedDocument(
            elements=elements,
            markdown=md,
            page_count=max(pages_seen) if pages_seen else 0,
            table_count=table_count,
            parser_used="unstructured",
        )

    # -------------------- pdfplumber (本地) --------------------

    async def _parse_with_pdfplumber(self, path: Path) -> ParsedDocument:
        import pdfplumber

        elements: List[Element] = []
        table_count = 0

        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                table_areas = []
                if tables:
                    for tbl_obj in page.find_tables() or []:
                        table_areas.append(tbl_obj.bbox)

                text = page.extract_text() or ""
                if text.strip():
                    for line in text.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        etype = _guess_element_type(line)
                        elements.append(Element(
                            type=etype,
                            text=line,
                            metadata={"page_number": page_num},
                        ))

                for tbl in tables:
                    table_count += 1
                    md_table = _raw_table_to_markdown(tbl)
                    elements.append(Element(
                        type=ElementType.TABLE,
                        text=md_table,
                        metadata={"page_number": page_num},
                    ))

        md = _elements_to_markdown(elements)
        return ParsedDocument(
            elements=elements,
            markdown=md,
            page_count=page_count,
            table_count=table_count,
            parser_used="pdfplumber",
        )

    # -------------------- PyPDF2 (兜底) --------------------

    async def _parse_with_pypdf2(self, path: Path) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                logger.error("pypdf 和 PyPDF2 均未安装，无法解析 PDF")
                return ParsedDocument(parser_used="none")

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        elements: List[Element] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    elements.append(Element(
                        type=_guess_element_type(line),
                        text=line,
                        metadata={"page_number": page_num},
                    ))

        md = _elements_to_markdown(elements)
        return ParsedDocument(
            elements=elements,
            markdown=md,
            page_count=page_count,
            table_count=0,
            parser_used="pypdf2",
        )

    # -------------------- DOCX --------------------

    async def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 未安装，无法解析 DOCX")
            return ParsedDocument(parser_used="none")

        doc = Document(str(path))
        elements: List[Element] = []
        table_count = 0

        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p" and child in para_map:
                para = para_map[child]
                text = para.text.strip()
                if not text:
                    continue
                style = (para.style.name or "").lower() if para.style else ""
                if "heading" in style:
                    elements.append(Element(type=ElementType.TITLE, text=text))
                elif "list" in style:
                    elements.append(Element(type=ElementType.LIST_ITEM, text=text))
                else:
                    elements.append(Element(type=ElementType.NARRATIVE_TEXT, text=text))
            elif tag == "tbl" and child in table_map:
                table_count += 1
                md = _docx_table_to_markdown(table_map[child])
                if md:
                    elements.append(Element(type=ElementType.TABLE, text=md))

        md = _elements_to_markdown(elements)
        page_est = max(1, len(doc.paragraphs) // 30)
        return ParsedDocument(
            elements=elements,
            markdown=md,
            page_count=page_est,
            table_count=table_count,
            parser_used="python-docx",
        )


# ==================== 辅助函数 ====================


def _html_table_to_markdown(html: str) -> Optional[str]:
    """将 Unstructured 返回的 HTML 表格转为 Markdown"""
    try:
        import re
        rows_raw = re.findall(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL)
        if not rows_raw:
            return None
        md_rows = []
        for row_html in rows_raw:
            cells = re.findall(r"<t[hd][^>]*>(.*?)</t[hd]>", row_html, re.DOTALL)
            cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
            md_rows.append("| " + " | ".join(cells) + " |")
        if len(md_rows) < 1:
            return None
        col_count = md_rows[0].count("|") - 1
        separator = "| " + " | ".join(["---"] * max(col_count, 1)) + " |"
        return "\n".join([md_rows[0], separator] + md_rows[1:])
    except Exception:
        return None


def _map_unstructured_type(raw_type: str) -> ElementType:
    mapping = {
        "Title": ElementType.TITLE,
        "NarrativeText": ElementType.NARRATIVE_TEXT,
        "Table": ElementType.TABLE,
        "ListItem": ElementType.LIST_ITEM,
        "Header": ElementType.HEADER,
        "Footer": ElementType.FOOTER,
        "PageBreak": ElementType.PAGE_BREAK,
        "Image": ElementType.IMAGE,
        "Formula": ElementType.FORMULA,
    }
    return mapping.get(raw_type, ElementType.UNCATEGORIZED)


def _guess_element_type(line: str) -> ElementType:
    """从文本行猜测元素类型（用于 pdfplumber/PyPDF2 降级场景）"""
    import re
    stripped = line.strip()
    if re.match(r"^(\d+\.)+\s+\S", stripped) or re.match(r"^[A-Z][A-Z\s]{5,}$", stripped):
        return ElementType.TITLE
    if re.match(r"^[-•●◦▪]\s", stripped):
        return ElementType.LIST_ITEM
    return ElementType.NARRATIVE_TEXT


def _elements_to_markdown(elements: List[Element]) -> str:
    """将 Element 列表转换为 Markdown 文本"""
    parts: list[str] = []
    for el in elements:
        if el.type == ElementType.TITLE:
            parts.append(f"\n## {el.text}\n")
        elif el.type == ElementType.TABLE:
            parts.append(f"\n{el.text}\n")
        elif el.type == ElementType.LIST_ITEM:
            parts.append(f"- {el.text}")
        else:
            parts.append(el.text)
    return "\n".join(parts)


def _raw_table_to_markdown(rows: List[List]) -> str:
    """将 pdfplumber 提取的原始表格转换为 Markdown"""
    if not rows:
        return ""
    md_rows = []
    for row in rows:
        cells = [str(c or "").strip().replace("\n", " ") for c in row]
        md_rows.append("| " + " | ".join(cells) + " |")
    if len(md_rows) < 1:
        return ""
    col_count = len(rows[0])
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    return "\n".join([md_rows[0], separator] + md_rows[1:])


def _docx_table_to_markdown(table) -> str:
    """将 python-docx Table 对象转换为 Markdown 表格"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) < 1:
        return ""
    col_count = len(table.rows[0].cells) if table.rows else 0
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    return "\n".join([rows[0], separator] + rows[1:])


# ==================== 单例 ====================


_default_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    global _default_parser
    if _default_parser is None:
        _default_parser = DocumentParser()
    return _default_parser
