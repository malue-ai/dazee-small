"""
SAP 文档组装（Pandoc 主转换 + python-docx 后处理）

流程：
  1. 按 ASSEMBLY_ORDER 将各章节 Markdown 拼接为完整 .md
  2. 去重：按章节标题检测重复内容，仅保留最后一次生成的版本
  3. 清理 Markdown 残留（**bold** -> 保留原文，--- 移除）
  4. Pandoc 用 sap_reference.docx 模板转为 Word
  5. python-docx 后处理：[AI-INFERRED] 黄色高亮、[PLACEHOLDER] 红色高亮
  6. 无 Pandoc 时自动降级到纯 python-docx 方案

用法：python assemble_docx.py <chapters_dir> <entities_path> <output_path>
"""
import json, os, re, shutil, subprocess, sys, tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

FALLBACK_ASSEMBLY_ORDER = [
    ("00_title_page.md", "Title Page"),
    ("01_estimand.md", "1. Objectives, Endpoints, and Estimands"),
    ("02_study_design.md", "2. Study Design"),
    ("03_multiplicity.md", "3. Statistical Hypotheses and Multiplicity Adjustment"),
    ("04_analysis_sets.md", "4. Analysis Sets"),
    ("05_1_general.md", "5.1 General Considerations"),
    ("05_2_primary.md", "5.2 Primary Endpoint Analysis"),
    ("05_3_secondary.md", "5.3 Secondary Endpoint Analysis"),
    ("05_4_other_secondary.md", "5.4 Other Secondary and Exploratory Analyses"),
    ("05_5_safety.md", "5.5 Safety Analyses"),
    ("05_6_subgroup.md", "5.6 Other Analyses"),
    ("05_7_interim_changes.md", "5.7 Interim Analysis / Changes from Protocol"),
    ("06_sample_size.md", "6. Sample Size Determination"),
    ("07_references.md", "7. References"),
    ("08_appendix_abbreviations.md", "Appendix A: Abbreviations"),
]

SCRIPT_DIR = Path(__file__).parent
REFERENCE_DOCX = SCRIPT_DIR.parent / "styles" / "sap_reference.docx"


def _load_assembly_order(chapters_dir: Path):
    """Load assembly order from template_structure.json if available, else fallback."""
    ts_path = chapters_dir.parent / "template_structure.json"
    if not ts_path.exists():
        ts_path = chapters_dir / ".." / "template_structure.json"
    if ts_path.exists():
        try:
            data = json.loads(ts_path.read_text(encoding="utf-8"))
            sections = data.get("sections", [])
            if sections:
                order = []
                for i, sec in enumerate(sections):
                    sid = sec.get("id", str(i + 1))
                    title = sec.get("title", f"Section {sid}")
                    safe_id = re.sub(r"[^a-zA-Z0-9_.-]", "_", sid)
                    fname = f"{i:02d}_{safe_id}.md"
                    order.append((fname, f"{sid} {title}" if sid[0].isdigit() else title))
                print(f"  Loaded {len(order)} sections from template_structure.json")
                return order
        except Exception as e:
            print(f"  WARN: Failed to load template_structure.json: {e}")
    return FALLBACK_ASSEMBLY_ORDER


# ==================== Step 1: 拼接 + 去重 ====================


def merge_chapters(chapters_dir: Path):
    """将各章节 .md 按顺序拼接，自动去重，返回 (合并文本, 已包含数, 缺失列表)。"""
    assembly_order = _load_assembly_order(chapters_dir)
    parts, count, missing = [], 0, []

    for fname, title in assembly_order:
        f = chapters_dir / fname
        if not f.exists():
            # scan for any file containing the section id in its name
            candidates = [p for p in chapters_dir.glob("*.md") if fname.split("_", 1)[-1].replace(".md", "") in p.stem]
            if candidates:
                f = candidates[0]
        if f.exists():
            content = f.read_text(encoding="utf-8")
            parts.append(f"# {title}\n\n{content}\n\n\\newpage\n")
            count += 1
            print(f"  + {title} ({f.stat().st_size} bytes)")
        else:
            missing.append(fname)

    merged = "\n".join(parts)
    merged = _dedup_sections(merged)
    merged = _clean_markdown_residuals(merged)
    return merged, count, missing


def _dedup_sections(text: str) -> str:
    """Heading-level dedup: if the same heading appears multiple times, keep only the last occurrence."""
    heading_re = re.compile(r"^(#{1,3}\s+.+)$", re.MULTILINE)
    headings = heading_re.findall(text)

    seen = {}
    for h in headings:
        normalized = h.strip().lower()
        seen.setdefault(normalized, []).append(h)

    duplicates = {k: v for k, v in seen.items() if len(v) > 1}
    if not duplicates:
        return text

    for norm_key, occurrences in duplicates.items():
        to_remove = occurrences[:-1]
        for heading in to_remove:
            pattern = re.escape(heading) + r"\n(.*?)(?=\n#{1,3}\s|\n\\newpage|\Z)"
            match = re.search(pattern, text, re.DOTALL)
            if match:
                text = text[:match.start()] + text[match.end():]
                print(f"  DEDUP: removed duplicate section: {heading.strip()}")

    return text


def _clean_markdown_residuals(text: str) -> str:
    """Remove Markdown artifacts that should not appear in the final document."""
    text = re.sub(r"^---\s*$", "", text, flags=re.MULTILINE)
    return text


# ==================== Step 2: Pandoc 转换 ====================


def _pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def convert_with_pandoc(md_text: str, output_path: str) -> bool:
    """用 Pandoc + reference.docx 将 Markdown 转为 Word，成功返回 True。"""
    if not _pandoc_available():
        print("WARN: pandoc 未安装，降级到 python-docx 方案")
        return False

    with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(md_text)
        tmp_md = f.name

    try:
        cmd = [
            "pandoc", tmp_md,
            "-o", output_path,
            "--from", "markdown+pipe_tables+grid_tables+header_attributes",
            "--to", "docx",
            "--toc", "--toc-depth=3",
        ]
        if REFERENCE_DOCX.exists():
            cmd += ["--reference-doc", str(REFERENCE_DOCX)]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            print(f"WARN: pandoc 失败 (rc={result.returncode}): {result.stderr[:500]}")
            return False
        return True
    except Exception as e:
        print(f"WARN: pandoc 异常: {e}")
        return False
    finally:
        os.unlink(tmp_md)


# ==================== Step 3: python-docx 后处理 ====================


MARKER_PATTERN = re.compile(r"(\[AI-INFERRED\]|\[PLACEHOLDER[^\]]*\])")


def postprocess_markers(docx_path: str) -> None:
    """扫描 Word 文档中的 [AI-INFERRED]/[PLACEHOLDER] 文本，添加高亮。"""
    if not DOCX_OK:
        return
    doc = Document(docx_path)
    changed = False

    for para in doc.paragraphs:
        changed |= _highlight_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    changed |= _highlight_paragraph(para)

    if changed:
        doc.save(docx_path)
        print("  POST: [AI-INFERRED]/[PLACEHOLDER] 标记已高亮")


def _highlight_paragraph(para) -> bool:
    """对段落中包含标记的 run 拆分并高亮，返回是否有修改。"""
    full_text = para.text
    if not MARKER_PATTERN.search(full_text):
        return False

    for run in list(para.runs):
        run_text = run.text
        if not MARKER_PATTERN.search(run_text):
            continue

        parts = MARKER_PATTERN.split(run_text)
        if len(parts) <= 1:
            continue

        parent = run._element.getparent()
        run_idx = list(parent).index(run._element)

        run.text = parts[0]

        for k, part in enumerate(parts[1:], 1):
            if not part:
                continue
            from docx.oxml.ns import qn as _qn
            from copy import deepcopy
            new_r = deepcopy(run._element)
            new_r.text = ""
            rpr = new_r.find(_qn("w:rPr"))
            if rpr is None:
                rpr = new_r.makeelement(_qn("w:rPr"), {})
                new_r.insert(0, rpr)
            t = new_r.makeelement(_qn("w:t"), {})
            t.text = part
            t.set(_qn("xml:space"), "preserve")
            new_r.append(t)

            if part == "[AI-INFERRED]":
                hl = rpr.makeelement(_qn("w:highlight"), {_qn("w:val"): "yellow"})
                rpr.append(hl)
            elif part.startswith("[PLACEHOLDER"):
                hl = rpr.makeelement(_qn("w:highlight"), {_qn("w:val"): "red"})
                rpr.append(hl)

            parent.insert(run_idx + k, new_r)

    return True


# ==================== Fallback: 纯 python-docx ====================


def _parse_md_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        cells = [c.strip() for c in stripped.split("|")]
        if cells and all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def _add_word_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                tcpr = cell._element.get_or_add_tcPr()
                shd = tcpr.makeelement(qn("w:shd"), {qn("w:fill"): "D9E2F3", qn("w:val"): "clear"})
                tcpr.append(shd)


def _add_run_with_markers(para, text):
    parts = MARKER_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part)
        if part == "[AI-INFERRED]":
            run.font.highlight_color = 7
        elif part.startswith("[PLACEHOLDER"):
            run.font.highlight_color = 6


def convert_with_python_docx(md_text: str, output_path: str) -> None:
    """纯 python-docx 降级方案。"""
    if not DOCX_OK:
        print("ERROR: pip install python-docx")
        sys.exit(1)

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.54)
        s.left_margin = s.right_margin = Cm(3.17)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("\\newpage"):
            doc.add_page_break()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.lstrip().startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].rstrip().lstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            _add_word_table(doc, _parse_md_table(table_lines))
            continue
        else:
            p = doc.add_paragraph()
            _add_run_with_markers(p, line)
        i += 1

    doc.save(output_path)


# ==================== Main ====================


def main(chapters_dir: str, entities_path: str, output_path: str) -> None:
    chapters = Path(chapters_dir)
    print(f"Assembling SAP from: {chapters}")

    md_text, count, missing = merge_chapters(chapters)
    if count == 0:
        print("ERROR: 没有找到任何章节文件")
        sys.exit(1)

    if convert_with_pandoc(md_text, output_path):
        print(f"  PANDOC: {count} chapters -> {output_path}")
        postprocess_markers(output_path)
    else:
        print("  FALLBACK: python-docx 方案")
        convert_with_python_docx(md_text, output_path)

    print(f"\nAssembled {count}/{len(ASSEMBLY_ORDER)} chapters -> {output_path}")
    if missing:
        print(f"Missing: {missing}")


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print(f"Usage: {sys.argv[0]} <chapters_dir> <entities> <output>")
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])
