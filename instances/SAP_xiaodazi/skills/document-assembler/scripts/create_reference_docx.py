"""
生成 Pandoc reference.docx 模板（SAP 法规文档样式）

Pandoc 通过 --reference-doc 参数继承模板中的 Word 样式定义。
此脚本创建一个符合 ICH E9/EU-PEARL SAP 规范的 reference.docx：
  - 字体：正文 Times New Roman 11pt，标题 Arial Bold
  - 页边距：上下 2.54cm，左右 3.17cm
  - 表格：Table Grid + 表头浅蓝底色 + 9pt
  - 页眉页脚：Confidential 标记 + 页码

用法：python create_reference_docx.py [output_path]
默认输出到同级 styles/sap_reference.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


def create_reference(output_path: str) -> None:
    doc = Document()

    # --- 页面设置 ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = "CONFIDENTIAL"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = "Arial"

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run().font.size = Pt(8)

    # --- 样式定义 ---
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style_normal.element.rPr
    if rpr is None:
        rpr = style_normal.element.makeelement(qn("w:rPr"), {})
        style_normal.element.append(rpr)
    rpr_ea = rpr.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Times New Roman"})
    rpr.append(rpr_ea)

    pf = style_normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level, size, space_before in [(1, 14, 24), (2, 12, 18), (3, 11, 12)]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.name = "Arial"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)
            hs.paragraph_format.space_before = Pt(space_before)
            hs.paragraph_format.space_after = Pt(6)

    # Pandoc 的表格样式需要名为 "Table" 的自定义样式；
    # 由于 python-docx 不支持直接创建表格样式，Pandoc 会回退到 Table Grid。
    # 这里写入一个示例表格让 Pandoc 学习格式。

    # --- 示例内容（Pandoc 从此继承样式映射）---
    doc.add_heading("Heading 1 Example", level=1)
    doc.add_heading("Heading 2 Example", level=2)
    doc.add_heading("Heading 3 Example", level=3)
    doc.add_paragraph("Body text paragraph in Times New Roman 11pt.")
    doc.add_paragraph("Bullet item example", style="List Bullet")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for j, header_text in enumerate(["Column A", "Column B", "Column C"]):
        cell = table.rows[0].cells[j]
        cell.text = header_text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D9E2F3",
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)
    for j, val in enumerate(["Data 1", "Data 2", "Data 3"]):
        cell = table.rows[1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

    doc.save(output_path)
    print(f"Reference template saved: {output_path}")


if __name__ == "__main__":
    default_out = str(Path(__file__).parent.parent / "styles" / "sap_reference.docx")
    out = sys.argv[1] if len(sys.argv) > 1 else default_out
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    create_reference(out)
